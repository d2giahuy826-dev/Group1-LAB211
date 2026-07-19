from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import csv, json

root=Path(__file__).parent; project=root.parent.parent
summary=list(csv.DictReader((root/'experiment_summary.csv').open(encoding='utf-8')))
rows=list(csv.DictReader((root/'experiment_results.csv').open(encoding='utf-8')))
doc=Document(); sec=doc.sections[0]
sec.page_height=Inches(11); sec.page_width=Inches(8.5)
sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
styles=doc.styles
normal=styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(11)
normal.paragraph_format.space_after=Pt(8); normal.paragraph_format.line_spacing=1.22
for name,size,color,before,after in [('Title',28,'17324D',0,10),('Heading 1',16,'2E74B5',16,8),('Heading 2',13,'2E74B5',12,6),('Heading 3',12,'1F4D78',8,4)]:
    s=styles[name]; s.font.name='Calibri'; s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color)
    s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)
header=sec.header.paragraphs[0]; header.text='LAB211  |  PAYROLL CONCURRENCY RESEARCH'; header.style=styles['Caption']
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT
footer.add_run('Nhóm 1  •  '); fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); footer._p.append(fld)

def page(title, kicker=None):
    if len(doc.paragraphs)>1: doc.add_page_break()
    if kicker:
        p=doc.add_paragraph(); r=p.add_run(kicker.upper()); r.bold=True; r.font.size=Pt(9); r.font.color.rgb=RGBColor(54,162,105)
    doc.add_heading(title,0)
def h(text,l=1): doc.add_heading(text,level=l)
def p(text,boldlead=None):
    q=doc.add_paragraph()
    if boldlead: q.add_run(boldlead).bold=True
    q.add_run(text); return q
def bullets(items):
    for item in items: doc.add_paragraph(item,style='List Bullet')
def table(headers,data,widths=None):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'; t.autofit=False
    for i,x in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=str(x); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for r in c.paragraphs[0].runs: r.bold=True; r.font.color.rgb=RGBColor(255,255,255)
        shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),'2F6B8A'); c._tc.get_or_add_tcPr().append(shd)
    for row in data:
        cells=t.add_row().cells
        for i,x in enumerate(row): cells[i].text=str(x)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    return t
def fig(path,caption,width=6.3):
    doc.add_picture(str(path),width=Inches(width)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    c=doc.add_paragraph(caption); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; c.style='Caption'

page('Nghiên cứu cơ chế đồng bộ trong hệ thống Payroll','Báo cáo thực nghiệm Task 9')
p('Đánh giá NO_LOCK, SYNCHRONIZED, OPTIMISTIC và FILE_LOCK dưới tải 20 HR threads.')
doc.add_paragraph('\n'); q=doc.add_paragraph('BÁO CÁO KỸ THUẬT • 2026'); q.alignment=WD_ALIGN_PARAGRAPH.CENTER
q=doc.add_paragraph('Group 1 — LAB211, FPT University'); q.alignment=WD_ALIGN_PARAGRAPH.CENTER
p('\nDữ liệu: 20 phòng ban × 50 nhân viên = 1.000 nhân viên; mỗi cơ chế chạy độc lập 3 lần.')

page('Tóm tắt điều hành','01 • Executive summary')
p('Thực nghiệm cho thấy khóa loại bỏ hoàn toàn thanh toán trùng, nhưng dữ liệu phép nền đang không nhất quán với đơn đã duyệt. Vì vậy không có cơ chế nào đạt 0% tổng lỗi trong trạng thái dữ liệu hiện tại.')
h('Kết quả nổi bật')
bullets(['NO_LOCK nhanh nhất (10.733 emp/s) nhưng có 95% duplicate trên số lượt ghi thành công.',
         'SYNCHRONIZED nhanh nhất trong nhóm không tạo double payment: 4.046 emp/s.',
         'OPTIMISTIC đạt 3.639 emp/s và 0 duplicate; phù hợp khi tranh chấp thấp.',
         'FILE_LOCK chỉ đạt 36,6 emp/s do chi phí mở/khóa file cho từng entry.',
         'Cả bốn cơ chế phát hiện 1.775 sai lệch phép, tương đương 88,75% trên 2.000 phép đối chiếu.'])
h('Trả lời Research Question')
p('Không có cơ chế đạt đồng thời 0% lỗi và nhanh nhất vì wrongLeave là lỗi dữ liệu nền, không thay đổi theo cơ chế. Nếu chỉ xét lỗi cạnh tranh double payment, SYNCHRONIZED là cơ chế 0 lỗi nhanh nhất trong mẫu đo này.')

page('Mục lục','02 • Navigation')
table(['Phần','Nội dung','Trang'],[['1–3','Bối cảnh, câu hỏi, phạm vi','4–6'],['4–6','Kiến trúc MVC và thiết kế đồng bộ','7–9'],['7–9','Phương pháp và dữ liệu','10–12'],['10–12','Kết quả và phân tích','13–15'],['13–15','Kết luận, hạn chế, phụ lục','16–18']],[1.0,4.5,1.0])
h('Danh mục hiện vật')
bullets(['Bảng 1 — Biến độc lập, biến phụ thuộc và điều kiện kiểm soát','Hình 1 — Class diagram hệ thống','Hình 2 — Throughput trung bình','Hình 3 — Tỷ lệ lỗi hậu kiểm','Hình 4 — Đánh đổi lỗi–hiệu năng','Bảng 4 — Toàn bộ 12 lần chạy'])

page('Bối cảnh và vấn đề nghiên cứu','03 • Introduction')
p('Payroll là miền nghiệp vụ nhạy cảm: một race condition có thể tạo thanh toán lặp hoặc trừ phép sai. Kiến trúc lưu CSV làm cửa sổ read–modify–write rõ ràng hơn và cho phép quan sát trực tiếp hậu quả cạnh tranh.')
h('Vấn đề')
bullets(['Nhiều HR thread xử lý cùng một payroll entry.', 'Kiểm tra “đã xử lý” và ghi CSV không nguyên tử trong NO_LOCK.', 'Khóa quá rộng bảo đảm đúng nhưng có thể làm giảm throughput.', 'Dữ liệu phép cần được hậu kiểm độc lập với bộ đếm trong RAM.'])
h('Mục tiêu')
p('Đo lường định lượng mức độ đúng đắn và hiệu năng của bốn chiến lược, sau đó lựa chọn cơ chế dựa trên bằng chứng thay vì suy đoán.')

page('Câu hỏi nghiên cứu và giả thuyết','04 • Research design')
h('Research Question')
p('Trong bốn cơ chế NO_LOCK, SYNCHRONIZED, OPTIMISTIC và FILE_LOCK, cơ chế nào đạt 0% lỗi và nhanh nhất khi 20 HR threads xử lý 1.000 nhân viên?')
h('Giả thuyết trước thực nghiệm')
table(['Mã','Giả thuyết'],[['H1','NO_LOCK có throughput cao nhất nhưng phát sinh duplicate.'],['H2','Ba cơ chế khóa không phát sinh duplicate.'],['H3','OPTIMISTIC nhanh hơn SYNCHRONIZED khi xung đột được loại sớm.'],['H4','FILE_LOCK chậm nhất do chi phí khóa ở tầng hệ điều hành.']],[.7,5.8])
h('Tiêu chí quyết định')
p('Ưu tiên tuyệt đối: tổng lỗi bằng 0. Nếu nhiều cơ chế cùng đạt 0, chọn TPS trung bình cao nhất. Nếu không cơ chế nào đạt 0, báo cáo nguyên nhân và không thay đổi tiêu chí sau khi xem kết quả.')

page('Phạm vi, biến và chỉ số','05 • Measurement')
table(['Loại','Biến / chỉ số','Định nghĩa'],[['Độc lập','Lock mechanism','4 mức: NO_LOCK, SYNCHRONIZED, OPTIMISTIC, FILE_LOCK'],['Kiểm soát','Workload','20 threads; 20 departments; 50 employees/dept'],['Phụ thuộc','elapsedMs','Thời gian wall-clock cho một lần chạy'],['Phụ thuộc','TPS','success × 1.000 / elapsedMs'],['Chất lượng','doublePayment','Số entryId dư sau khi đọc lại CSV'],['Chất lượng','wrongLeave','Số annualUsed/sickUsed lệch tổng ngày APPROVED']],[1.0,1.6,3.9])
h('Đơn vị phân tích')
p('Mỗi hàng dữ liệu là một lần chạy độc lập. Trung bình được tính trên ba lần chạy của cùng cơ chế. Detector luôn mở lại file CSV sau khi simulation kết thúc; không sử dụng bộ đếm lỗi trong RAM.')

page('Kiến trúc MVC của hệ thống','06 • Architecture')
p('MainView chỉ điều phối menu; SimulatorController quản lý use case; PayrollSimulationService sở hữu concurrency logic; repositories và detector thực hiện I/O/hậu kiểm. Việc tách lớp giữ business logic khỏi View và làm test độc lập.')
fig(project/'docs'/'class_diagram.png','Hình 1. Class diagram tổng thể của dự án.',5.9)
h('Luồng dữ liệu')
p('Input CSV → Controller → ExecutorService/CountDownLatch → cơ chế đồng bộ → output CSV riêng → Detector → SimulationResult → View/ExperimentRunner.')

page('Thiết kế bộ mô phỏng','07 • Simulator internals')
h('Khởi động đồng thời')
p('ExecutorService tạo đúng số worker. CountDownLatch ready bảo đảm tất cả worker đã sẵn sàng; latch start mở cùng thời điểm; latch done chặn controller đến khi hoàn tất.')
h('Đảm bảo tái lập')
bullets(['Mỗi cơ chế ghi file output riêng.', 'Danh sách payroll entry lấy từ cùng source CSV.', 'Kết quả thô giữ đủ 12 dòng và thông số tải.', 'Mỗi detector đọc lại CSV sau khi worker hoàn tất.'])
h('Đo thời gian')
p('System.nanoTime() bao quanh toàn bộ lệnh service.run; elapsed được đổi sang milliseconds và chặn tối thiểu 1 ms để tránh chia cho 0.')

page('Bốn cơ chế đồng bộ','08 • Concurrency strategies')
table(['Cơ chế','Điểm nguyên tử','Đánh đổi'],[['NO_LOCK','Không có','Nhanh nhưng nhiều thread cùng ghi một entry.'],['SYNCHRONIZED','Monitor JVM + processedIds','Đơn giản, đúng trong một process.'],['OPTIMISTIC','Token entryId@version','Reject sớm khi version đã được claim.'],['FILE_LOCK','FileChannel.lock()','Bảo vệ tầng file, chi phí I/O/OS cao.']],[1.25,2.1,3.15])
h('Lưu ý diễn giải')
p('Tập processedIds loại bỏ phép quét O(n²) khỏi đường nóng của SYNCHRONIZED và FILE_LOCK. FILE_LOCK vẫn khóa file thật khi ghi. Đây là tối ưu thuật toán kiểm tra, không loại bỏ cơ chế đang được đo.')

page('Detector và tính toàn vẹn dữ liệu','09 • Post-simulation audit')
h('Double payment')
p('Detector nhóm các dòng theo entryId và cộng max(0, count−1). Với 1.000 entry và 20 threads, NO_LOCK tạo 20.000 lượt ghi thành công, trong đó 19.000 là bản sao dư.')
h('Wrong leave deduction')
p('Với từng LeaveBalance, detector so annualUsed với tổng days của đơn ANNUAL APPROVED và sickUsed với tổng days của đơn SICK APPROVED. Có 2.000 phép đối chiếu cho 1.000 nhân viên.')
h('Phát hiện dữ liệu nền')
p('Sai lệch 1.775 xuất hiện giống nhau ở mọi cơ chế. Đây là bằng chứng lỗi nằm ở dữ liệu sinh ban đầu/luồng cập nhật phép, không phải race condition của payroll simulation.')

page('Quy trình thực nghiệm','10 • Protocol')
bullets(['Biên dịch Java 17 và chạy unit test.', 'Chạy ExperimentRunner với 20 threads, 20 departments, 50 employees/dept.', 'Lặp ba lần theo thứ tự NO_LOCK → SYNCHRONIZED → OPTIMISTIC → FILE_LOCK.', 'Xuất experiment_results.csv và experiment_results.json.', 'Tính trung bình theo cơ chế và tạo ba biểu đồ.', 'Đối chiếu kết luận với tiêu chí quyết định đã định trước.'])
h('Mối đe dọa tới độ tin cậy')
bullets(['Thứ tự chạy cố định có thể chịu warm-up/JIT.', 'Windows file-system cache ảnh hưởng FILE_LOCK.', 'Ba lần lặp đủ cho bài tập nhưng chưa đủ cho benchmark công nghiệp.', 'wrongLeave dùng snapshot dữ liệu nền thay vì workload cập nhật phép đồng thời.'])

page('Kết quả throughput','11 • Performance results')
fig(root/'throughput_comparison.png','Hình 2. TPS trung bình trên ba lần chạy.',6.3)
table(['Cơ chế','elapsed TB (ms)','TPS TB'],[[s['mechanism'],f"{float(s['avgElapsedMs']):,.1f}",f"{float(s['avgTPS']):,.1f}"] for s in summary],[2.2,2.0,2.0])
p('NO_LOCK dẫn đầu tuyệt đối nhưng con số success gồm cả thanh toán lặp. Trong nhóm loại bỏ duplicate, SYNCHRONIZED nhanh hơn OPTIMISTIC khoảng 11,2%; FILE_LOCK thấp hơn SYNCHRONIZED hơn 110 lần.')

page('Kết quả tỷ lệ lỗi','12 • Correctness results')
fig(root/'error_rates.png','Hình 3. Tỷ lệ double payment và wrong leave hậu kiểm.',6.3)
p('NO_LOCK có double-payment rate 95%. Ba cơ chế còn lại đạt 0% duplicate. Wrong-leave rate 88,75% không đổi, xác nhận đây là bất nhất snapshot dữ liệu chứ không phải hiệu ứng của payroll lock.')
h('Không đánh đồng success với correct success')
p('20.000 “success” của NO_LOCK không có nghĩa xử lý được 20.000 nhân viên; chỉ có 1.000 entry duy nhất. 19.000 dòng dư phải được xem là lỗi tài chính nghiêm trọng.')

page('Phân tích đánh đổi','13 • Trade-off')
fig(root/'tradeoff_scatter.png','Hình 4. Vị trí của từng cơ chế theo throughput và tổng tỷ lệ lỗi quan sát.',6.1)
h('Diễn giải')
p('Khi cộng hai tỷ lệ kiểm tra, không điểm nào nằm trên trục 0% lỗi. NO_LOCK ở vùng nhanh–rủi ro cao; SYNCHRONIZED và OPTIMISTIC ở vùng duplicate-safe nhưng còn lỗi dữ liệu phép; FILE_LOCK vừa còn lỗi nền vừa chậm.')

page('Trả lời Research Question','14 • Decision')
h('Kết luận theo đúng tiêu chí')
p('Không có cơ chế nào đạt 0% tổng lỗi trong full experiment. Do đó không thể tuyên bố một “cơ chế 0 lỗi nhanh nhất” nếu không che giấu 1.775 sai lệch phép.')
h('Quyết định có điều kiện')
table(['Phạm vi quyết định','Khuyến nghị','Căn cứ'],[['Chỉ payroll duplicate','SYNCHRONIZED','0 duplicate; 4.045,6 emp/s — nhanh nhất nhóm an toàn.'],['Tranh chấp thấp, cần versioning','OPTIMISTIC','0 duplicate; 3.638,7 emp/s; conflict rõ ràng.'],['Nhiều process ghi chung file','FILE_LOCK','0 duplicate nhưng chỉ 36,6 emp/s.'],['Toàn hệ thống','Chưa nghiệm thu','Phải sửa pipeline LeaveBalance trước.']],[1.55,1.6,3.35])

page('Khuyến nghị kỹ thuật','15 • Recommendations')
bullets(['Chọn SYNCHRONIZED cho bản console một JVM hiện tại.', 'Giữ OPTIMISTIC cho hướng phát triển database/multi-service.', 'Không dùng NO_LOCK cho dữ liệu tài chính.', 'Chỉ dùng FILE_LOCK khi thật sự có nhiều process và không thể chuyển sang DB transaction.', 'Sửa DataGenerator để annualUsed/sickUsed được suy ra từ leave_requests APPROVED.', 'Thêm invariant test chạy sau sinh dữ liệu: wrongLeave phải bằng 0.', 'Tách successAttempt khỏi uniqueProcessed trong KPI để không thưởng cho duplicate.'])
h('Acceptance gate đề xuất')
p('Một build chỉ được demo khi doublePayment=0, wrongLeave=0, số entry duy nhất=1.000 và tất cả test JUnit thành công.')

page('Hạn chế và hướng nghiên cứu tiếp','16 • Limitations')
bullets(['Đo trên CSV local, chưa đại diện database transaction.', 'Chưa random hóa thứ tự cơ chế giữa các lần chạy.', 'Không đo p95 latency hoặc CPU/memory.', 'Workload leave chưa chạy đồng thời trong cùng experiment.', 'File lock đo mỗi entry; batching có thể cải thiện mạnh throughput.'])
h('Thực nghiệm tiếp theo')
p('Chạy 10–30 lần với warm-up; random hóa thứ tự; bổ sung mức thread 1/5/10/20/50; đo latency distribution; thử batch file lock; đưa dữ liệu vào database và so pessimistic/optimistic transaction.')

page('Phụ lục A — 12 lần chạy','17 • Raw evidence')
table(['Mechanism','Run','Success','Duplicate','Wrong leave','ms','TPS'],[[r['mechanism'],r['run'],r['success'],r['doublePayment'],r['wrongLeave'],r['elapsedMs'],f"{float(r['TPS']):.1f}"] for r in rows],[1.3,.45,.75,.85,.85,.65,.75])
p('Nguồn: experiment_results.csv được ExperimentRunner sinh trực tiếp sau full experiment ngày 17/07/2026.')

page('Phụ lục B — Tái lập và kết luận','18 • Reproducibility')
h('Tệp tái lập')
bullets(['src/simulation/ExperimentRunner.java — orchestration 4 × 3.', 'docs/experiment/experiment_results.csv — dữ liệu thô.', 'docs/experiment/experiment_results.json — bản JSON.', 'docs/experiment/experiment_summary.csv — số trung bình.', 'src/test/*Simulation*Test.java — regression tests.'])
h('Kết luận cuối')
p('Thực nghiệm chứng minh race condition không khóa là lỗi có thể tái hiện, không phải giả thuyết. Đồng thời, hậu kiểm độc lập đã phát hiện một vấn đề quan trọng hơn phạm vi payroll: dữ liệu phép nền không thỏa invariant. Thành công của Task 9 nằm ở việc trả lời trung thực bằng số liệu và biến detector thành quality gate cho các vòng phát triển tiếp theo.')

doc.save(project/'docs'/'report.docx')
